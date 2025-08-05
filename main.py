from fastapi import FastAPI
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt
import uuid
import base64
import os

app = FastAPI()

class ResumeRequest(BaseModel):
    content: str

@app.post("/format-resume")
def format_resume(req: ResumeRequest):
    content = req.content
    output_filename = f"resume_{uuid.uuid4().hex}.docx"
    output_path = f"/tmp/{output_filename}"

    # Create the formatted resume
    document = Document()
    style = document.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    lines = content.split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            document.add_paragraph()
            continue

        paragraph = document.add_paragraph()
        run = paragraph.add_run()

        if line.startswith("**") and line.endswith("**"):
            run.text = line.strip("*")
            run.bold = True
            run.font.size = Pt(12)
        elif line.startswith("*") and line.endswith("*"):
            run.text = line.strip("*")
            run.italic = True
        elif line.startswith("- "):
            run.text = line
            paragraph.paragraph_format.left_indent = Pt(15)
        else:
            run.text = line

    document.save(output_path)

    # Encode the file in base64 and return
    with open(output_path, "rb") as file:
        encoded = base64.b64encode(file.read()).decode("utf-8")

    return {
        "filename": output_filename,
        "file_base64": encoded
    }
